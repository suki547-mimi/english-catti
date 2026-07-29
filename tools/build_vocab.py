"""Build normalized vocabulary JSON from local sources.

Input:  C:\\Cursorworkspace\\English\\Files\\*.doc / *.docx
Output: C:\\Cursorworkspace\\English\\data\\vocab.json
        C:\\Cursorworkspace\\English\\data\\corpus_sentences.json  (bilingual sentences from white paper)
        C:\\Cursorworkspace\\English\\data\\build_report.json      (per-topic stats)

Format of one vocab entry:
{
  "id": "sha1-of-normalized-en+zh",
  "zh": "安居工程",
  "en": "housing project for low-income families",
  "headword": "housing",           # first significant English word, lowercased
  "letter": "H",                    # A-Z index tab
  "topic": "01-社会发展",           # from source file
  "sources": ["01"],                # source topic codes (deduped)
  "kind": "phrase" | "word" | "sentence"  # rough classification
}
"""
from __future__ import annotations
import os, re, json, glob, hashlib, unicodedata
from typing import List, Dict, Tuple, Optional

ROOT   = r"C:\Cursorworkspace\English"
FILES  = os.path.join(ROOT, "Files")
OUTDIR = os.path.join(ROOT, "data")

# ---------------------------------------------------------------------------
# 1. Extract raw paragraphs from every source file (Word COM for .doc)
# ---------------------------------------------------------------------------
_word_app = None
def _word():
    global _word_app
    if _word_app is None:
        import pythoncom
        pythoncom.CoInitialize()
        from win32com.client import DispatchEx
        _word_app = DispatchEx("Word.Application")
        _word_app.Visible = False
        _word_app.DisplayAlerts = 0
    return _word_app

def extract_docx(path: str) -> List[str]:
    from docx import Document
    doc = Document(path)
    out = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                out.append(cell.text)
    return out

def extract_doc(path: str) -> List[str]:
    tmp = path + ".tmp.docx"
    w = _word()
    d = w.Documents.Open(os.path.abspath(path), ReadOnly=True, AddToRecentFiles=False)
    try:
        d.SaveAs2(os.path.abspath(tmp), FileFormat=16)  # wdFormatXMLDocument
    finally:
        d.Close(SaveChanges=0)
    try:
        return extract_docx(tmp)
    finally:
        try: os.remove(tmp)
        except OSError: pass

# ---------------------------------------------------------------------------
# 2. Line-level cleanup
# ---------------------------------------------------------------------------
RE_HAS_ZH   = re.compile(r"[\u4e00-\u9fff]")
RE_HAS_EN   = re.compile(r"[A-Za-z]")
RE_ONLY_ZH  = re.compile(r"^[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\s\d\-–—、，。：；()（）《》【】]*$")
RE_ONLY_EN  = re.compile(r"^[A-Za-z0-9\s\-\.,;:'\"()/&%\?\!\+\*]*$")

# noise patterns to drop entire lines
NOISE_PATTERNS = [
    re.compile(r"^口译笔译分类词汇"),                       # section header re-appearing
    re.compile(r"更多资料加微博"),
    re.compile(r"翻硕之家"),
    re.compile(r"QQ\d{5,}"),
    re.compile(r"^www\.TopSage\.com"),
    re.compile(r"^\s*$"),
]

def is_noise(line: str) -> bool:
    return any(p.search(line) for p in NOISE_PATTERNS)

def normalize_ws(s: str) -> str:
    # collapse runs of tabs / non-breaking / multiple spaces
    s = s.replace("\u00a0", " ").replace("\u3000", " ")
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()

# Some source paragraphs jam many entries together separated by \n or tab-tab.
def explode_line(line: str) -> List[str]:
    # First split on embedded newlines (docx paragraphs can contain \n for line breaks).
    parts: List[str] = []
    for chunk in line.split("\n"):
        chunk = chunk.strip()
        if not chunk:
            continue
        # Multi-column table-like layout: 2+ consecutive tabs might indicate a column break.
        # But keep single tabs, they're likely ZH<->EN separators inside one entry.
        parts.extend(re.split(r"\t{2,}", chunk))
    return [normalize_ws(p) for p in parts if p.strip()]

# ---------------------------------------------------------------------------
# 3. Parse a single cleaned line into (zh, en) pair
# ---------------------------------------------------------------------------
# Boundary detection: find the best split by counting Chinese vs Latin runs.
def _is_zh(c: str) -> bool: return "\u4e00" <= c <= "\u9fff"
def _is_en(c: str) -> bool: return c.isascii() and c.isalpha()

def _best_split(s: str) -> Optional[int]:
    """Return the split index that best separates a ZH cluster from an EN cluster.

    Strategy:
      * Locate the last ZH char and the last EN char.
      * If last ZH < last EN  -> line ends in Latin, treat as ZH-first. Split just
        AFTER the last ZH char, so trailing acronyms embedded in Chinese phrases
        (e.g. "AA制") stay with the Chinese side.
      * If last EN < last ZH  -> line ends in Chinese, treat as EN-first. Split
        just BEFORE the first ZH char.
      * If neither side has both scripts after the split, return None.
    """
    last_zh = -1
    last_en = -1
    first_zh = -1
    for i, c in enumerate(s):
        if _is_zh(c):
            last_zh = i
            if first_zh == -1: first_zh = i
        elif _is_en(c):
            last_en = i
    if last_zh == -1 or last_en == -1:
        return None
    if last_zh < last_en:
        # ZH-first: split just after the last ZH char.
        return last_zh + 1
    else:
        # EN-first: split just before the first ZH char.
        return first_zh

def parse_pair(line: str) -> Optional[Tuple[str, str]]:
    """Return (zh, en) or None if line isn't a bilingual entry."""
    line = normalize_ws(line)
    if not line or is_noise(line):
        return None
    has_zh = bool(RE_HAS_ZH.search(line))
    has_en = bool(RE_HAS_EN.search(line))
    if not (has_zh and has_en):
        return None
    # Fast path: explicit tab separator anywhere.
    if "\t" in line:
        left, _, right = line.partition("\t")
        left, right = left.strip(), right.strip()
        if RE_HAS_ZH.search(left) and RE_HAS_EN.search(right):
            return _clean_pair(left, right)
        if RE_HAS_EN.search(left) and RE_HAS_ZH.search(right):
            return _clean_pair(right, left)
    # Otherwise find best split.
    flip = _best_split(line)
    if flip is None:
        return None
    left, right = line[:flip].strip(), line[flip:].strip()
    if not left or not right:
        return None
    if RE_HAS_ZH.search(left) and RE_HAS_EN.search(right):
        return _clean_pair(left, right)
    if RE_HAS_EN.search(left) and RE_HAS_ZH.search(right):
        return _clean_pair(right, left)
    return None

_LEAD_TRAIL_PUNCT = ",.;:!?，。；：？！\"'()（）[]【】《》〈〉 -–—"

def _clean_pair(zh: str, en: str) -> Tuple[str, str]:
    zh = zh.strip().strip(_LEAD_TRAIL_PUNCT)
    en = en.strip().strip(_LEAD_TRAIL_PUNCT)
    # Kill leading list markers "1、", "1)", "(1)"
    zh = re.sub(r"^\s*\d+\s*[、\.\)）]\s*", "", zh)
    en = re.sub(r"^\s*\d+\s*[\.\)]\s*", "", en)
    zh = normalize_ws(zh)
    en = normalize_ws(en)
    return zh, en

# ---------------------------------------------------------------------------
# 4. Headword + classification
# ---------------------------------------------------------------------------
STOP_HEAD = {"a", "an", "the", "to", "of", "in", "on", "for", "with", "by", "at",
             "and", "or", "not", "no", "as", "be", "is", "are", "was", "were"}

def extract_headword(en: str) -> str:
    """Pick the first content word, or fall back to first token."""
    toks = re.findall(r"[A-Za-z][A-Za-z\-']+", en)
    for t in toks:
        if t.lower() not in STOP_HEAD:
            return t.lower()
    return toks[0].lower() if toks else ""

def classify(zh: str, en: str) -> str:
    words = en.split()
    if len(words) >= 8 or re.search(r"[.!?。！？]", en):
        return "sentence"
    if len(words) == 1 and len(zh) <= 8:
        return "word"
    return "phrase"

def make_id(zh: str, en: str) -> str:
    key = f"{zh.lower()}|||{en.lower()}"
    return hashlib.sha1(key.encode("utf-8")).hexdigest()[:16]

# ---------------------------------------------------------------------------
# 5. Topic code from filename
# ---------------------------------------------------------------------------
TOPIC_MAP = {
    "01": "01-社会发展", "02": "02-外经贸", "03": "03-改革开放",
    "04": "04-政治政府", "05": "05-环境保护", "06": "06-机关机构",
    "07": "07-世贸组织", "08": "08-经济金融", "09": "09-热门话题",
    "10": "10-文化教育", "11": "11-宗教事务", "12": "12-商务英语",
    "13": "13-常见职位", "14": "14-电脑网络", "15": "15-科学技术",
    "16": "16-新闻经典", "17": "17-常用称谓", "19": "19-谚语",
    "20": "20-常用口语",
}

def topic_of(filename: str) -> Optional[str]:
    m = re.search(r"分类词汇（(\d+)）", filename)
    return TOPIC_MAP.get(m.group(1)) if m else None

# ---------------------------------------------------------------------------
# 6. White paper -> bilingual sentence corpus (for Q5 cloze)
# ---------------------------------------------------------------------------
def build_sentence_corpus(paragraphs: List[str]) -> List[Dict]:
    """White paper alternates ZH paragraph then EN paragraph. Pair them up."""
    cleaned = [p.strip() for p in paragraphs if p.strip()]
    pairs = []
    i = 0
    while i < len(cleaned) - 1:
        a, b = cleaned[i], cleaned[i + 1]
        a_zh = bool(RE_HAS_ZH.search(a)) and not bool(RE_HAS_EN.search(a[:20]))
        b_en = bool(RE_HAS_EN.search(b)) and not bool(RE_HAS_ZH.search(b))
        if a_zh and b_en:
            pairs.append({"zh": a, "en": b, "source": "peacekeeping-white-paper-2020"})
            i += 2
        else:
            i += 1
    return pairs

# ---------------------------------------------------------------------------
# 7. Main pipeline
# ---------------------------------------------------------------------------
def main():
    os.makedirs(OUTDIR, exist_ok=True)

    # Group topical files, keeping ONE representative per topic code (dedup 19x3).
    doc_paths = sorted(glob.glob(os.path.join(FILES, "口译笔译分类词汇*.doc")))
    seen_topics = set()
    picked = []
    for p in doc_paths:
        m = re.search(r"分类词汇（(\d+)）", os.path.basename(p))
        if not m:
            continue
        code = m.group(1)
        if code in seen_topics:
            continue
        seen_topics.add(code)
        picked.append(p)

    entries: Dict[str, Dict] = {}   # id -> entry (accumulates sources)
    per_topic_stats: Dict[str, Dict[str, int]] = {}

    for p in picked:
        name = os.path.basename(p)
        topic = topic_of(name)
        print(f"[extract] {name}  (topic={topic})")
        raw_paras = extract_doc(p)
        # Explode jammed paragraphs (05, 12) on \n
        lines: List[str] = []
        for para in raw_paras:
            lines.extend(explode_line(para))

        added = 0
        skipped = 0
        for line in lines:
            pair = parse_pair(line)
            if pair is None:
                skipped += 1
                continue
            zh, en = pair
            # Post-filters: reject junk entries
            if len(zh) < 2 or len(en) < 2:
                skipped += 1; continue
            if len(zh) > 300 or len(en) > 500:
                skipped += 1; continue

            eid = make_id(zh, en)
            if eid in entries:
                if topic and topic not in entries[eid]["sources"]:
                    entries[eid]["sources"].append(topic)
                continue
            headword = extract_headword(en)
            entries[eid] = {
                "id": eid,
                "zh": zh,
                "en": en,
                "headword": headword,
                "letter": headword[:1].upper() if headword else "#",
                "topic": topic or "misc",
                "sources": [topic] if topic else [],
                "kind": classify(zh, en),
            }
            added += 1

        per_topic_stats[topic or name] = {
            "raw_lines": len(lines),
            "added": added,
            "skipped": skipped,
        }

    # White paper -> sentence corpus (not into vocab)
    wp = os.path.join(FILES, "202009《中国军队参加联合国维和行动30年》白皮书双语全文.docx")
    sentences = []
    if os.path.exists(wp):
        print(f"[extract] {os.path.basename(wp)}  (sentence corpus)")
        sentences = build_sentence_corpus(extract_docx(wp))

    if _word_app is not None:
        try: _word_app.Quit()
        except Exception: pass

    # Sort vocab by letter, then headword, then zh
    vocab = sorted(entries.values(), key=lambda e: (e["letter"], e["headword"], e["zh"]))

    with open(os.path.join(OUTDIR, "vocab.json"), "w", encoding="utf-8") as f:
        json.dump(vocab, f, ensure_ascii=False, indent=1)
    with open(os.path.join(OUTDIR, "corpus_sentences.json"), "w", encoding="utf-8") as f:
        json.dump(sentences, f, ensure_ascii=False, indent=1)

    # Report
    letter_counts: Dict[str, int] = {}
    topic_counts: Dict[str, int] = {}
    kind_counts: Dict[str, int] = {}
    for e in vocab:
        letter_counts[e["letter"]] = letter_counts.get(e["letter"], 0) + 1
        topic_counts[e["topic"]]   = topic_counts.get(e["topic"], 0) + 1
        kind_counts[e["kind"]]     = kind_counts.get(e["kind"], 0) + 1

    report = {
        "vocab_total": len(vocab),
        "sentences_total": len(sentences),
        "per_source": per_topic_stats,
        "by_letter": dict(sorted(letter_counts.items())),
        "by_topic":  dict(sorted(topic_counts.items())),
        "by_kind":   dict(sorted(kind_counts.items())),
    }
    with open(os.path.join(OUTDIR, "build_report.json"), "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n[done] vocab={len(vocab)}  sentences={len(sentences)}")
    print(f"       -> {OUTDIR}\\vocab.json / corpus_sentences.json / build_report.json")

if __name__ == "__main__":
    main()
