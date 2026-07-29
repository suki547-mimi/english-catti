"""Quick sampling to judge whether each source file is actually useful
as a CATTI vocabulary source. For every file:
  - report size, page/paragraph count
  - detect if it contains structured "English word / Chinese meaning" pairs
  - print head & tail snippets
Runs on Windows with Word COM for legacy .doc.
"""
from __future__ import annotations
import os, re, sys, glob, json, io

ROOT = r"C:\Cursorworkspace\English\Files"
OUT  = r"C:\Cursorworkspace\English\tools\sample_report.json"

# --- extractors -------------------------------------------------------------
def extract_docx(path):
    from docx import Document
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                tables.append(" | ".join(cells))
    return paras, tables

def extract_pdf(path):
    import fitz
    d = fitz.open(path)
    paras = []
    for page in d:
        txt = page.get_text("text")
        for line in txt.splitlines():
            line = line.strip()
            if line:
                paras.append(line)
    return paras, []

_word_app = None
def _word():
    global _word_app
    if _word_app is None:
        import pythoncom
        pythoncom.CoInitialize()
        from win32com.client import DispatchEx
        _word_app = DispatchEx("Word.Application")
        _word_app.Visible = False
        _word_app.DisplayAlerts = 0
    return _word_app

def extract_doc(path):
    # Use Word COM to convert to a temp .docx then read with python-docx
    tmp = path + ".tmp.docx"
    w = _word()
    doc = w.Documents.Open(os.path.abspath(path), ReadOnly=True, AddToRecentFiles=False)
    try:
        # 16 = wdFormatXMLDocument (docx)
        doc.SaveAs2(os.path.abspath(tmp), FileFormat=16)
    finally:
        doc.Close(SaveChanges=0)
    try:
        return extract_docx(tmp)
    finally:
        try: os.remove(tmp)
        except OSError: pass

# --- analysis ---------------------------------------------------------------
RE_EN_ZH = re.compile(r"^([A-Za-z][A-Za-z\-\s\.,/&']{1,60})[\s\t]{2,}([\u4e00-\u9fff][^A-Za-z]{0,80})$")
RE_HAS_EN = re.compile(r"[A-Za-z]{3,}")
RE_HAS_ZH = re.compile(r"[\u4e00-\u9fff]")

def analyse(lines):
    n = len(lines)
    en_lines = sum(1 for l in lines if RE_HAS_EN.search(l))
    zh_lines = sum(1 for l in lines if RE_HAS_ZH.search(l))
    mix      = sum(1 for l in lines if RE_HAS_EN.search(l) and RE_HAS_ZH.search(l))
    pairs = [l for l in lines if RE_EN_ZH.match(l)]
    return {
        "lines": n,
        "lines_with_en": en_lines,
        "lines_with_zh": zh_lines,
        "lines_mixed_en_zh": mix,
        "pair_like": len(pairs),
        "pair_ratio": round(len(pairs) / n, 3) if n else 0.0,
    }

def head_tail(lines, k=8):
    return {"head": lines[:k], "tail": lines[-k:] if len(lines) > k else []}

# --- main -------------------------------------------------------------------
def main():
    files = []
    for ext in ("*.doc", "*.docx", "*.pdf"):
        files.extend(sorted(glob.glob(os.path.join(ROOT, ext))))

    report = []
    for i, path in enumerate(files, 1):
        name = os.path.basename(path)
        size_kb = round(os.path.getsize(path) / 1024, 1)
        print(f"[{i}/{len(files)}] {name}  ({size_kb} KB)")
        try:
            ext = os.path.splitext(path)[1].lower()
            if ext == ".docx":
                paras, tables = extract_docx(path)
            elif ext == ".pdf":
                paras, tables = extract_pdf(path)
            elif ext == ".doc":
                paras, tables = extract_doc(path)
            else:
                continue
            all_lines = paras + tables
            entry = {
                "file": name,
                "size_kb": size_kb,
                "ext": ext,
                "paragraphs": len(paras),
                "table_rows": len(tables),
                **analyse(all_lines),
                **head_tail(all_lines),
            }
        except Exception as e:
            entry = {"file": name, "size_kb": size_kb, "error": f"{type(e).__name__}: {e}"}
        report.append(entry)

    # Clean up Word
    if _word_app is not None:
        try: _word_app.Quit()
        except Exception: pass

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    with open(OUT, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"\n=> wrote {OUT}  ({len(report)} files)")

if __name__ == "__main__":
    main()
